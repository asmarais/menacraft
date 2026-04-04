import pdfplumber
import docx
import re
import unicodedata
import io
import numpy as np
from langdetect import detect as detect_language
import spacy
from sentence_transformers import SentenceTransformer

# Load models once
_nlp = spacy.load("en_core_web_sm")
_sentence_model = SentenceTransformer("all-MiniLM-L6-v2")  # free, local


def preprocess_document(doc_input: bytes, filename: str) -> dict:
    """
    Master document preprocessor
    """
    ext = filename.lower().split(".")[-1]

    # ── STAGE 1: Extract raw text + metadata ─────────────────────
    if ext == "pdf":
        raw_text, doc_meta, layout_features = extract_from_pdf(doc_input)
    elif ext in ("docx", "doc"):
        raw_text, doc_meta, layout_features = extract_from_docx(doc_input)
    else:  # txt / plain
        raw_text = doc_input.decode("utf-8", errors="ignore")
        doc_meta = {}
        layout_features = {}

    # ── STAGE 2+3: Text cleaning & NLP ───────────────────────────
    text_features = extract_text_features(raw_text)

    # ── STAGE 5: Embeddings ───────────────────────────────────────
    embedding_features = extract_embedding_features(raw_text)

    return {
        "input_type": "document",
        "filename": filename,
        "extension": ext,

        # ── For Axis 1 (Authenticity)
        "authenticity_features": {
            "ai_text_model_input": text_features["clean_text"][:2000],
            "burstiness_ratio": text_features["burstiness_ratio"],
            "sentence_length_variance": text_features["sentence_length_variance"],
            "avg_sentence_length": text_features["avg_sentence_length"],
            "perplexity_proxy": text_features["perplexity_proxy"],
            "metadata_anomalies": doc_meta.get("anomalies", []),
            "layout_anomalies": layout_features.get("anomalies", []),
            "font_consistency_score": layout_features.get("font_consistency", 1.0),
        },

        # ── For Axis 2 (Contextual Consistency)
        "context_features": {
            "clean_text": text_features["clean_text"],
            "sentence_embeddings": embedding_features["sentence_embeddings"],
            "document_embedding": embedding_features["document_embedding"],
            "entities": text_features["entities"],
            "keywords": text_features["keywords"],
            "topic_signals": text_features["topic_signals"],
            "language": text_features["language"],
            "sentiment": text_features["sentiment"],
        },

        # ── For Axis 3 (Source Credibility)
        "source_features": {
            "author": doc_meta.get("author", ""),
            "creator_tool": doc_meta.get("creator", ""),
            "creation_date": doc_meta.get("creation_date", ""),
            "modification_date": doc_meta.get("modification_date", ""),
            "has_digital_signature": doc_meta.get("has_signature", False),
            "page_count": doc_meta.get("page_count", 1),
            "word_count": text_features["word_count"],
            "writing_style_signals": text_features["style_signals"],
            "urls_found": text_features["urls_found"],
        }
    }


def extract_from_pdf(pdf_bytes: bytes) -> tuple[str, dict, dict]:
    full_text = []
    fonts_used = set()
    font_sizes = []
    images_found = 0
    anomalies = []

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        page_count = len(pdf.pages)
        meta = pdf.metadata or {}

        for page in pdf.pages:
            # Text
            text = page.extract_text() or ""
            full_text.append(text)

            # Font analysis
            for char in (page.chars or []):
                fonts_used.add(char.get("fontname", ""))
                font_sizes.append(char.get("size", 12))

            # Images
            images_found += len(page.images or [])

        # Font consistency
        font_size_std = float(np.std(font_sizes)) if font_sizes else 0.0
        if len(fonts_used) > 5:
            anomalies.append("excessive_font_variety")
        if font_size_std > 8:
            anomalies.append("inconsistent_font_sizes")

    doc_meta = {
        "author": meta.get("Author", "") or meta.get("/Author", ""),
        "creator": meta.get("Creator", "") or meta.get("/Creator", ""),
        "producer": meta.get("Producer", "") or meta.get("/Producer", ""),
        "creation_date": meta.get("CreationDate", ""),
        "modification_date": meta.get("ModDate", ""),
        "page_count": page_count,
        "has_signature": False,  # would need PyMuPDF for this
        "anomalies": anomalies,
    }

    layout_features = {
        "fonts_used": list(fonts_used),
        "font_count": len(fonts_used),
        "font_consistency": max(0.0, 1.0 - len(fonts_used) / 10),
        "has_embedded_images": images_found > 0,
        "image_count": images_found,
        "anomalies": anomalies,
    }

    return "\n".join(full_text), doc_meta, layout_features


def extract_from_docx(docx_bytes: bytes) -> tuple[str, dict, dict]:
    doc = docx.Document(io.BytesIO(docx_bytes))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    core_props = doc.core_properties
    doc_meta = {
        "author": core_props.author or "",
        "creator": "Microsoft Word / LibreOffice",
        "creation_date": str(core_props.created or ""),
        "modification_date": str(core_props.modified or ""),
        "page_count": len(doc.sections),
        "has_signature": False,
        "anomalies": [],
    }

    fonts_in_doc = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name:
                fonts_in_doc.add(run.font.name)

    layout_features = {
        "fonts_used": list(fonts_in_doc),
        "font_count": len(fonts_in_doc),
        "font_consistency": max(0.0, 1.0 - len(fonts_in_doc) / 8),
        "anomalies": [],
    }

    return full_text, doc_meta, layout_features


def extract_text_features(raw_text: str) -> dict:
    # Clean
    text = unicodedata.normalize("NFKC", raw_text)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\s+', ' ', text).strip()

    # Basic stats
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 5]
    lengths = [len(s.split()) for s in sentences]

    word_count = len(text.split())
    avg_len = float(np.mean(lengths)) if lengths else 0
    std_len = float(np.std(lengths)) if lengths else 0
    burstiness = std_len / (avg_len + 1e-9)
    perplexity_proxy = std_len  # higher = more human-like variation

    # Language
    try:
        lang = detect_language(text[:500])
    except:
        lang = "unknown"

    # NER
    doc_nlp = _nlp(text[:5000])
    entities = [
        {"text": ent.text, "label": ent.label_}
        for ent in doc_nlp.ents
    ]

    # Keywords (TF-IDF style via frequency)
    from collections import Counter
    stopwords = {"the","a","an","is","in","it","of","and","to","was","for","on","that"}
    words = [w.lower() for w in text.split() if w.isalpha() and w.lower() not in stopwords]
    keywords = [w for w, _ in Counter(words).most_common(15)]

    # URLs
    urls = re.findall(r'https?://\S+', text)

    # Sentiment (simple heuristic)
    positive_words = {"good","great","confirmed","verified","official","true"}
    negative_words = {"breaking","shocking","exclusive","unbelievable","secret","hidden"}
    lower_text = text.lower()
    sentiment = {
        "sensationalism_score": sum(1 for w in negative_words if w in lower_text) / len(negative_words),
        "positive_signal": sum(1 for w in positive_words if w in lower_text),
    }

    # Style signals
    style_signals = {
        "uses_caps_frequently": len(re.findall(r'\b[A-Z]{3,}\b', text)) > 3,
        "excessive_exclamations": text.count("!") > 5,
        "question_heavy": text.count("?") > 10,
        "all_uppercase_ratio": sum(1 for w in text.split() if w.isupper()) / max(word_count, 1),
    }

    # Topic signals
    political_terms = {"government","president","minister","election","protest","military"}
    crisis_terms = {"war","attack","bomb","killed","disaster","emergency"}
    topic_signals = {
        "political": sum(1 for t in political_terms if t in lower_text),
        "crisis": sum(1 for t in crisis_terms if t in lower_text),
    }

    return {
        "clean_text": text,
        "word_count": word_count,
        "sentence_count": len(sentences),
        "avg_sentence_length": round(avg_len, 2),
        "sentence_length_variance": round(std_len, 3),
        "burstiness_ratio": round(burstiness, 4),
        "perplexity_proxy": round(perplexity_proxy, 3),
        "language": lang,
        "entities": entities[:30],
        "keywords": keywords,
        "urls_found": urls,
        "sentiment": sentiment,
        "style_signals": style_signals,
        "topic_signals": topic_signals,
    }


def extract_embedding_features(text: str) -> dict:
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 10][:20]

    if not sentences:
        return {"sentence_embeddings": [], "document_embedding": []}

    embeddings = _sentence_model.encode(sentences, convert_to_numpy=True)
    doc_embedding = embeddings.mean(axis=0)

    return {
        "sentence_embeddings": embeddings.tolist(),
        "document_embedding": doc_embedding.tolist(),  # 384-dim vector
    }