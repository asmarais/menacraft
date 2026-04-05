import os
import shutil
import tempfile
import re
import asyncio
import logging
import requests
import httpx
from bs4 import BeautifulSoup
from typing import Optional
from fastapi import UploadFile
from app.axes.authenticity import Authenticity
from app.axes.credibility import Credibility
from app.axes.consistency import Consistency

logger = logging.getLogger("Pipeline")

class Pipeline:
    def __init__(self):
        self.auth = Authenticity()
        self.cred = Credibility()
        
        # Initialize Consistency with HF_TOKEN from env
        hf_token = os.environ.get("HF_TOKEN")
        self.consistency = Consistency(hf_token=hf_token)
        # Initialize temp directory
        self.temp_dir = os.path.join(os.getcwd(), "temp")
        os.makedirs(self.temp_dir, exist_ok=True)


    async def run(self, input_type: str, file: Optional[UploadFile] = None, url: Optional[str] = None, claim_text: Optional[str] = None) -> dict:
        """
        Coordinates preprocessing and axis evaluation for various input types.
        Runs Authenticity + Credibility axes in parallel.
        """
        features = {"type": input_type}
        temp_path = None

        try:
            # 1. ── Save File if provided ──
            if file:
                filename = f"upload_{os.urandom(4).hex()}_{file.filename}"
                temp_path = os.path.join(self.temp_dir, filename)
                with open(temp_path, "wb") as buffer:
                    shutil.copyfileobj(file.file, buffer)
                features["path"] = temp_path

            # 2. ── Enrich features (merge, never overwrite) ──
            features["url"] = url
            features["claim"] = claim_text or ""
            features["claim_text"] = claim_text or ""

            if temp_path:
                features["path"] = temp_path          # Authenticity._image reads "path"
                features["file_path"] = temp_path
                if input_type == "video":
                    features["video_path"] = temp_path
                elif input_type == "image":
                    features["image_path"] = temp_path
                elif input_type == "document":
                    features["document_path"] = temp_path

            # Specialized enrichment based on type
            if input_type == "url" and url:
                social_data = self._extract_social_metadata(url)
                if social_data:
                    features["account"] = social_data

                body_text = self._scrape_url(url)
                features["body_text"] = body_text
                features["clean_text"] = body_text   # for credibility _text fallback
                features["word_count"] = len(body_text.split())

                metrics = self._calculate_writing_metrics(body_text)
                features.update(metrics)

            elif input_type == "document" and temp_path:
                text_content = self._extract_document_text(temp_path)
                features["clean_text"] = text_content
                features["body_text"] = text_content  # for credibility _text fallback
                features["word_count"] = len(text_content.split())

                metrics = self._calculate_writing_metrics(text_content)
                features.update(metrics)

            # 3. ── Evaluate axes in parallel ──
            auth_task = self.auth.evaluate(features)
            cred_task = asyncio.to_thread(self.cred.evaluate, features)
            consistency_task = asyncio.to_thread(self.consistency.evaluate, features)

            auth_result, cred_result, consistency_result = await asyncio.gather(
                auth_task, cred_task, consistency_task, return_exceptions=True
            )

            if isinstance(auth_result, Exception):
                logger.error(f"Authenticity evaluation failed: {auth_result}")
                auth_result = {"score": 0.5, "label": "error", "explanation": str(auth_result), "flags": []}
            if isinstance(cred_result, Exception):
                logger.error(f"Credibility evaluation failed: {cred_result}")
                cred_result = {"score": 0.5, "label": "unknown", "explanation": str(cred_result), "flags": []}
            if isinstance(consistency_result, Exception):
                logger.error(f"Consistency evaluation failed: {consistency_result}")
                consistency_result = {"score": 0.5, "label": "unknown", "explanation": str(consistency_result), "flags": []}

            # --- TRANSFORMATION FOR FRONTEND ---
            # 1. Transform component_scores (dict -> array)
            cs_dict = auth_result.get("details", {}).get("component_scores", {})
            component_scores = []
            for name, score in cs_dict.items():
                if name != "ai_gen_by_source":
                    component_scores.append({
                        "name": name.replace("_", " ").title(),
                        "score": score,
                        "verdict": "fake" if score > 0.6 else "uncertain" if score > 0.3 else "real"
                    })

            # 2. Transform api_raw (dict -> array)
            api_dict = auth_result.get("details", {}).get("api_raw", {})
            api_raw = []
            for name, res in api_dict.items():
                if isinstance(res, dict):
                    score = res.get("ai_generated_score") or res.get("manipulation_score")
                    if score is not None:
                        api_raw.append({
                            "source": name.replace("_", " ").title(),
                            "score": float(score),
                            "verdict": str(res.get("verdict", "uncertain")).lower()
                        })

            # 3. Final Details
            details = auth_result.get("details", {}).copy()
            details["component_scores"] = component_scores
            details["api_raw"] = api_raw

            # ── Build Axes for frontend dashboard ──
            axes = [
                {
                    "axis": "Content Authenticity",
                    "score": auth_result.get("score", 0.5),
                    "verdict": auth_result.get("label", "uncertain"),
                    "flags": auth_result.get("flags", []),
                    "explanation": auth_result.get("explanation", "")
                }
            ]

            # ── Consistency Axis ──
            axes.append({
                "axis": "Contextual Consistency",
                "score": consistency_result.get("score", 0.5),
                "verdict": consistency_result.get("label", "unknown"),
                "flags": consistency_result.get("flags", []),
                "explanation": consistency_result.get("explanation", ""),
                "details": consistency_result.get("details", {})
            })

            # Individual Axis for Metadata (image/document)
            exif_score = cs_dict.get("exif_risk")
            if exif_score is not None:
                axes.append({
                    "axis": "Metadata Forensics",
                    "score": exif_score,
                    "verdict": "fake" if exif_score > 0.6 else "uncertain" if exif_score > 0.3 else "real",
                    "flags": [f for f in auth_result.get("flags", []) if "EXIF" in f or "Software" in f]
                })

            # ── Credibility Axis ──
            cred_score = cred_result.get("score", 0.5)
            cred_label = cred_result.get("label", "unknown")
            cred_flags = cred_result.get("flags", [])

            axes.append({
                "axis": "Source Credibility",
                "score": round(1.0 - cred_score, 4), # 1.0 = Low Credibility (Risk)
                "verdict": cred_label,
                "flags": cred_flags,
                "explanation": cred_result.get("explanation", ""),
                "details": cred_result.get("details", {})
            })

            # Add credibility details to the main details block
            details["credibility"] = {
                "score":       cred_score,
                "label":       cred_label,
                "explanation": cred_result.get("explanation", ""),
                "flags":       cred_flags,
            }

            # ── Compute combined verdict ──
            auth_score = auth_result.get("score", 0.5)
            const_score = consistency_result.get("score", 0.5)
            
            # Weighted: 50% Authenticity (1=Fake), 25% Credibility (0=Fake), 25% Consistency (1=Fake)
            combined_score = (auth_score * 0.50) + ((1.0 - cred_score) * 0.25) + (const_score * 0.25)
            combined_score = max(0.0, min(1.0, combined_score))
            
            # Additional details for Consistency in main block
            details["consistency"] = {
                "score":       const_score,
                "label":       consistency_result.get("label", "unknown"),
                "explanation": consistency_result.get("explanation", ""),
                "flags":       consistency_result.get("flags", []),
            }

            if combined_score < 0.35:
                combined_verdict = "real"
            elif combined_score < 0.50:
                combined_verdict = "uncertain"
            else:
                combined_verdict = "fake"

            # Use the authenticity explanation as primary (it's LLM-generated)
            explanation = auth_result.get("explanation", "")
            if cred_result.get("explanation") and cred_label in ("suspicious", "dangerous", "highly_suspicious"):
                explanation += f" Source credibility: {cred_result['explanation']}."

            return {
                "verdict": combined_verdict,
                "score": round(combined_score, 4),
                "explanation": explanation,
                "claim": claim_text,
                "axes": axes,
                "details": details,
                "type": input_type
            }

        finally:
            # Cleanup temp file? 
            # pass (Keep for now to debug)
            pass

    def _calculate_writing_metrics(self, text: str) -> dict:
        """
        Calculates linguistic metrics for burstiness analysis.
        """
        import numpy as np
        # Split by common sentence delimiters
        sentences = [s.strip() for s in re.split(r'[.!?]+', text) if len(s.strip()) > 10]
        
        if not sentences:
            return {
                "burstiness_ratio": 0.0,
                "sentence_length_variance": 0.0,
                "avg_sentence_length": 0.0
            }

        # Calculate word counts per sentence
        counts = [len(s.split()) for s in sentences]
        
        avg_len = np.mean(counts)
        std_len = np.std(counts)
        variance = np.var(counts)
        
        # Burstiness Ratio = StdDev / Mean (Uniformity coefficient)
        ratio = std_len / avg_len if avg_len > 0 else 0
        
        return {
            "burstiness_ratio": round(float(ratio), 4),
            "sentence_length_variance": round(float(variance), 4),
            "avg_sentence_length": round(float(avg_len), 4)
        }

    def _extract_document_text(self, file_path: str) -> str:
        """
        Extracts plain text from PDF or DOCX files.
        """
        ext = file_path.lower().split(".")[-1]
        text = ""
        try:
            if ext == "pdf":
                import pypdfium2 as pdfium
                pdf = pdfium.PdfDocument(file_path)
                for page in pdf:
                    text += page.get_textpage().get_text_range() + "\n"
            elif ext in ["docx", "doc"]:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([p.text for p in doc.paragraphs])
            else:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
        except Exception as e:
            text = f"Error extracting text: {str(e)}"
        
        return text.strip()

    def _scrape_url(self, url: str) -> str:
        """
        Scrapes article body text from a URL.
        """
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
                "Accept-Language": "en-US,en;q=0.9",
                "Referer": "https://www.google.com/"
            }
            with httpx.Client(follow_redirects=True, timeout=10.0) as client:
                response = client.get(url, headers=headers)
                response.raise_for_status()
                soup = BeautifulSoup(response.text, "html.parser")
            
            # Remove scripts, styles
            for s in soup(["script", "style", "nav", "footer"]): s.decompose()
            
            # Try to find common article tags
            article = soup.find("article") or soup.find("main") or soup
            text = article.get_text(separator=" ", strip=True)
            
            # Clean up whitespace
            text = re.sub(r"\s+", " ", text)
            return text[:8000] # Increased cap
        except Exception as e:
            logger.warning(f"Static scrape failed for {url}: {e}. This may happen if the site blocks crawlers.")
            return f"Scraping failed: {str(e)}"

    def _extract_social_metadata(self, url: str) -> Optional[dict]:
        """
        Attempts to extract social media profile metadata from HTML meta tags.
        """
        try:
            headers = {
                "User-Agent": "Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Referer": "https://www.google.com/"
            }
            with httpx.Client(follow_redirects=True, timeout=10.0) as client:
                response = client.get(url, headers=headers)
                html = response.text
                soup = BeautifulSoup(html, "html.parser")

            data = {
                "account_age_days": 730, # Default to 2 years for neutral
                "followers": 0,
                "following": 0,
                "posts_last_24h": 0,
                "verified": False,
                "bio": "",
                "profile_pic": True,
                "website": False,
                "total_posts": 0,
                "source": "meta_scrape"
            }

            # 1. ── Extract Meta/Bio ──
            meta_desc_tag = soup.find("meta", property="og:description") or soup.find("meta", attrs={"name": "description"})
            meta_desc = meta_desc_tag.get("content", "") if meta_desc_tag else ""
            data["bio"] = meta_desc

            # 2. ── Platform Specific Logic ──
            if "instagram.com" in url:
                # Instagram OG description/title can contain stats
                combined = meta_desc + " | " + (soup.title.string if soup.title else "")
                og_title_tag = soup.find("meta", property="og:title")
                if og_title_tag: combined += " | " + og_title_tag.get("content", "")

                matches = {
                    "followers": re.search(r"([\d,.]+)\s*(Followers|Abonnés)", combined, re.I),
                    "following": re.search(r"([\d,.]+)\s*(Following|Abonnements)", combined, re.I),
                    "posts": re.search(r"([\d,.]+)\s*(Posts|Publications)", combined, re.I)
                }
                
                for key, match in matches.items():
                    if match:
                        val = match.group(1).replace(",", "")
                        try:
                            # Handle K/M suffixes in value
                            multiplier = 1
                            if "K" in val.upper(): multiplier = 1000; val = val.upper().replace("K", "")
                            if "M" in val.upper(): multiplier = 1000000; val = val.upper().replace("M", "")
                            
                            num = int(float(val) * multiplier)
                            if key == "posts": data["total_posts"] = num
                            else: data[key] = num
                        except: pass
                
                # --- Fallback: broad search in raw HTML if still 0 ---
                if data["followers"] == 0:
                    f_match = re.search(r"(\d+[\d,.]*[KMB]?)\s*(Followers|Abonnés)", html, re.I)
                    if f_match:
                        val = f_match.group(1).replace(",", "")
                        try: data["followers"] = int(float(val.replace("K","000").replace("M","000000"))) # simple handle
                        except: pass
                if data["total_posts"] == 0:
                    p_match = re.search(r"(\d+[\d,.]*[KMB]?)\s*(Posts|Publications)", html, re.I)
                    if p_match:
                        val = p_match.group(1).replace(",", "")
                        try: data["total_posts"] = int(float(val.replace("K","000").replace("M","000000")))
                        except: pass

            # 3. ── Verification Detection ──
            # Look for common "verified" keywords in meta and title
            page_title = soup.title.string if soup.title else ""
            if any(x in (html + page_title) for x in ["Verified", "VerifiedAccount", "verified-badge"]):
                data["verified"] = True

            # 4. ── Bio cleanup ──
            data["bio"] = data["bio"][:255] if data["bio"] else "No bio available"
            
            return data
        except Exception as e:
            logger.error(f"Metadata extraction failed for {url}: {e}")
            return None